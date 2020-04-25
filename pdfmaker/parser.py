from collections import defaultdict
import json
from bs4 import BeautifulSoup
from docx import Document

with open("book.json", "rt") as file:
    html_pages = json.load(file)


bad_phrases = {
    "Best Authors",
    "Popular Series",
    "Vampire Books",
    "Read Free Online",
    "Free Books"
}

full_book = []
for page_num in range(1, 100):
    parsed_page = BeautifulSoup(html_pages[str(page_num)])
    para_text = [p.get_text() for p in parsed_page.find_all('p')]
    for p in para_text:
        if any(bp in p for bp in bad_phrases):
            print("Dropping para:\n", p)
        else:
            full_book.append(p)



# Parse into chapters
chapter_broken_book = defaultdict(lambda: [])
chap = None
for para in full_book:

    if para.strip().isdigit():
        chap = int(para.strip())
        print(chap)

    chapter_broken_book[chap].append(para)


book = Document()

chapter_paras = chapter_broken_book[None]
book.add_heading(chapter_paras[0], level=1)
for p in chapter_paras[1:]:
    book.add_paragraph(p)

chapters = sorted(k for k in chapter_broken_book if k)
for chap in chapters:
    chapter_paras = chapter_broken_book[chap]
    book.add_heading(chapter_paras[0], level=2)
    for p in chapter_paras[1:]:
        book.add_paragraph(p)

    

with open("book.rtf", "wt") as file:
    file.write("\n".join(full_book))