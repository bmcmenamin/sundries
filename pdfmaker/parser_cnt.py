from collections import defaultdict
import json
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

import re



scraped_file = "reaper.json"
output_file = "reaper.docx"

chapter_re = re.compile(r"[\s]*Chapter [\d]+", re.IGNORECASE)
def is_chapter_head(para):
    if chapter_re.match(para):
        return int(para.split()[1])

def to_center(para):
    return para == "* * *"

bad_phrases = {}





with open(scraped_file, "rt") as file:
    html_pages = json.load(file)

parsed_page = BeautifulSoup(html_pages["1"])

for idx, wordsect in enumerate(parsed_page.findAll("div", {"class": "wordsection"})):
    break

    for chap_content in parsed_page.findAll("div", {"class": "chapter-content"}):



    para_text = []
    for chap_content in parsed_page.findAll("div", {"id": "textToRead"}):
        text_block = "".join(chap_content.get_text().strip())
        text_block = text_block.replace("\n\n\n\n\n\t\t\n\t\t\t", "").replace(u'\xa0', u'\n')
        lines = [
            line.rstrip()
            for line in text_block.splitlines()
        ]

        para_text += lines

    for p in para_text:
        if any(bp in p for bp in bad_phrases):
            print("Dropping para:\n", p)
        else:
            full_book.append(p)

full_book = full_book[(149 - 1): (4088 - 1)]


with open("test.txt", "wt") as file:
    file.write("\n".join(full_book))


# Parse into chapters
chapter_broken_book = defaultdict(lambda: [])
chap = None
for para in full_book:

    possible_chap = is_chapter_head(para)
    if possible_chap:
        chap = possible_chap
        chapter_broken_book[chap].append(para.strip())
    else:        
        chapter_broken_book[chap].append(para)


# Write into Docx

book = Document()

# Start with pre-chapter 1
chapter_paras = chapter_broken_book[None]
book.add_heading(chapter_paras[0], level=1)
for p in chapter_paras[1:]:
    para = book.add_paragraph(p.strip())
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if to_center(p):
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
book.add_page_break()

# Go through chapters in order
chapters = sorted(k for k in chapter_broken_book if k)
for chap in chapters:
    chapter_paras = chapter_broken_book[chap]
    book.add_heading(chapter_paras[0], level=2)
    for p in chapter_paras[1:]:
        para = book.add_paragraph(p.strip())
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if to_center(p):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    book.add_page_break()

book.save(output_file)

