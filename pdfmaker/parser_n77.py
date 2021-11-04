from collections import defaultdict
import json
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

import re
from num2words import num2words





scraped_file = "heart_of_obsidian.json"
output_file = "heart_of_obsidian.docx"

chapter_re = re.compile(r"[\s]*Chapter [\d]+", re.IGNORECASE)
def is_chapter_head(para):
    if chapter_re.match(para):
        return int(para.split()[1])

def to_center(para):
    return para == "* * *"

bad_phrases = {}




scraped_file = "reaver.json"
output_file = "reaver.docx"

possible_chapters = {
    num2words(c, ordinal=False, lang='en', to='cardinal'): c
    for c in range(1, 100)
}
def is_chapter_head(para):
    key = para.lower().strip().replace(" ", "-")
    return possible_chapters.get(key)

def to_center(para):
    return para == "* * *"

bad_phrases = {}





with open(scraped_file, "rt") as file:
    html_pages = json.load(file)

max_pages = max(int(k) for k in html_pages)


full_book = []
for page_num in range(1, max_pages + 1):
    parsed_page = BeautifulSoup(html_pages[str(page_num)])

    para_text = []
    for chap_content in parsed_page.findAll("div", {"class": "chapter-content"}):
        lines = [
            line.rstrip()
            for p in chap_content.find_all('p')
            for line in p.get_text().strip().splitlines()
        ]
        para_text += [l for l in lines if l]

    for p in para_text:
        if any(bp in p for bp in bad_phrases):
            print("Dropping para:\n", p)
        else:
            full_book.append(p)


with open("test.txt", "wt") as file:
    file.write("\n".join(full_book))


# Parse into chapters
chapter_broken_book = defaultdict(lambda: [])
chap = None
for para in full_book:

    possible_chap = is_chapter_head(para)
    if possible_chap:
        chap = possible_chap
        chapter_broken_book[chap].append(para.strip())
    else:        
        chapter_broken_book[chap].append(para)


# Write into Docx

book = Document()

# Start with pre-chapter 1
chapter_paras = chapter_broken_book[None]
book.add_heading(chapter_paras[0], level=1)
for p in chapter_paras[1:]:
    para = book.add_paragraph(p.strip())
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if to_center(p):
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
book.add_page_break()

# Go through chapters in order
chapters = sorted(k for k in chapter_broken_book if k)
for chap in chapters:
    chapter_paras = chapter_broken_book[chap]
    book.add_heading(chapter_paras[0], level=2)
    for p in chapter_paras[1:]:
        para = book.add_paragraph(p.strip())
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if to_center(p):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    book.add_page_break()

book.save(output_file)

